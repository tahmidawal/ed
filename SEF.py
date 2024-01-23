import openpyxl
import pandas as pd
import docx
from docx import Document
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
import datetime
import streamlit as st
from docx.oxml.ns import qn


# ************ START OVERALL HELPER FUNCTIONS ************ #


# ************ CHECK IF FILE PATHS EXIST ************ #

def check_file_exists(file_path):
    """
    Check if a file exists at the given file path.

    Args:
        file_path (str): The path of the file to check.

    Returns:
        bool: True if the file exists, False otherwise.
    """
    return os.path.isfile(file_path)


# ************ DATA EXTRACTION FROM PANDAS DATAFRAMES FUNCTIONS ************ #
def extract_rows_by_name(df, last_name, first_name):
    """
    Extracts rows from a DataFrame based on last name and first name.

    Parameters:
    df (pandas.DataFrame): The DataFrame to extract rows from.
    last_name (str): The last name to filter rows by.
    first_name (str): The first name to filter rows by.

    Returns:
    pandas.DataFrame: The filtered DataFrame.
    """
    # Adjust the condition to include rows with specific names or rows where both names are 'all'
    mask = ((df['First Name'].str.lower() == first_name.lower()) & (df['Last Name'].str.lower() == last_name.lower())) | \
        ((df['First Name'].str.lower() == 'all') & (df['Last Name'].str.lower() == 'all')) | \
        ((df['First Name'].str.lower() == 'all') & (df['Last Name'] == '')) | \
        ((df['First Name'] == '') & (df['Last Name'].str.lower() == 'all'))

    return df[mask]


def get_cell_contents(spreadsheet_path: str, row_number: int, column_name: str) -> str:
    """
    Retrieves the contents of a cell in a spreadsheet.

    Args:
        spreadsheet_path (str): The path to the spreadsheet file.
        row_number (int): The row number of the cell.
        column_name (str): The name of the column containing the cell.

    Returns:
        str: The contents of the cell.

    Raises:
        ValueError: If the column name is not found in the spreadsheet.
    """
    row_number = row_number + 1
    # Open the spreadsheet
    wb = openpyxl.load_workbook(spreadsheet_path)
    sheet = wb.active
    column_index = None

    # Get the column index by column name
    for i, row in enumerate(sheet.iter_rows()):
        for j, cell in enumerate(row):
            if cell.value == column_name:
                column_index = j
                break
        if column_index is not None:
            break
    if column_index is None:
        raise ValueError(f"{column_name} not found in the spreadsheet")

    # Get the cell contents
    return sheet.cell(row=row_number, column=column_index+1).value

# ************ END DATA EXTRACTION FROM PANDAS DATAFRAMES FUNCTIONS ************ #


# ************ START PARAGRAPH INSERTION FUNCTIONS ************ #


def insert_paragraph_with_font_style(file_path, text, font_size, font_style, font_color, header=False, highlight=False):
    """
    Inserts a paragraph with specified font style into a Word document.

    Args:
        file_path (str): The file path of the Word document.
        text (str): The text to be inserted as a paragraph.
        font_size (int): The font size of the paragraph.
        font_style (str): The font style of the paragraph.
        font_color (tuple): The RGB color values of the font in the format (R, G, B).
        header (bool, optional): Specifies whether the paragraph should be formatted as a header. Defaults to False.
        highlight (str, optional): Specifies the highlight color of the font. Can be "Blue", "Yellow", "Green", or "Red". Defaults to False.

    Returns:
        None
    """
    doc = docx.Document(file_path)
    paragraph = doc.add_paragraph(text)
    run = paragraph.runs[0]
    run.font.size = docx.shared.Pt(font_size)
    run.font.name = font_style
    run.font.color.rgb = docx.shared.RGBColor(
        font_color[0], font_color[1], font_color[2])
    if highlight == "Blue":
        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.BLUE
    elif highlight == "Yellow":
        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
    elif highlight == "Green":
        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GREEN
    elif highlight == "Red":
        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.RED
    if header:
        run.bold = True
    doc.save(file_path)


def copy_text_with_design_from_word_doc(source_file, destination_file):
    """
    Copy text with design from a Word document to another Word document.

    Args:
        source_file (str): The path of the source Word document.
        destination_file (str): The path of the destination Word document.

    Returns:
        None
    """
    print(source_file)
    src_doc = Document(source_file)
    dest_doc = Document(destination_file)

    for para in src_doc.paragraphs:
        new_para = dest_doc.add_paragraph()
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run._element.rPr.strike is not None:
                new_run_rPr = OxmlElement('w:rPr')
                new_strike = OxmlElement('w:strike')
                new_strike.set(qn('w:val'), 'true')
                new_run_rPr.append(new_strike)
                new_run._element.append(new_run_rPr)
            if run.font.subscript:
                new_run.sub()
            if run.font.superscript:
                new_run.font.superscript = True
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb
    dest_doc.save(destination_file)


def create_bulleted_list(word_file_path, items: list, start_index=None, font_size=None, font_color=None, font_style=None):
    """
    Create a bulleted list in a Word document.

    Args:
        word_file_path (str): The file path of the Word document.
        items (list): The list of items to be added as bullet points.
        start_index (int, optional): The index at which to start applying the bullet style. Defaults to None.
        font_size (int, optional): The font size of the bullet points. Defaults to None.
        font_color (str, optional): The font color of the bullet points. Defaults to None.
        font_style (str, optional): The font style of the bullet points. Defaults to None.
    """
    print(items)
    print(word_file_path)
    doc = docx.Document(word_file_path)
    for i in range(len(items)):
        print(items[i] + ": items[(i)]")
        run = doc.add_paragraph().add_run(items[i])
        if font_size is not None:
            run.font.size = docx.shared.Pt(font_size)
        if font_color is not None:
            run.font.color.rgb = docx.shared.RGBColor.from_string(font_color)
        if font_style is not None:
            run.font.name = font_style
        if start_index is not None:
            if i >= start_index:
                doc.paragraphs[-1].style = 'List Bullet'
        else:
            doc.paragraphs[-1].style = 'List Bullet'
    doc.save(word_file_path)


def create_numbered_list(client_file_path, general_items_list):
    """
    Creates a Word document with a numbered list from a given Python list.

    Args:
    general_items_list (list): A list of strings to be added to the Word document as a numbered list.
    client_file_path (str): The name of the Word file to be created.
    """
    # Create a new Document
    doc = docx.Document(client_file_path)

    # Loop through each item in the list and add it as a new paragraph
    for item in general_items_list:
        # Add paragraph with 'List Number' style for numbering
        doc.add_paragraph(item, style='List Number')

    # Save the document
    doc.save(client_file_path)


# ************ END PARAGRAPH INSERTION FUNCTIONS ************ #


# ************ START TABLE INSERTION FUNCTIONS ************ #


def bold_first_row(file_path, table_number):
    """
    Bold the first row of a table in a Word document.

    Args:
        file_path (str): The path to the Word document.
        table_number (int): The index of the table in the document.

    Returns:
        None
    """
    document = Document(file_path)
    table = document.tables[table_number]
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        # Change the text to white
    document.save(file_path)


def make_bold(file_path, row, column, table_number):
    """
    Makes the text in a specific cell of a table bold in a Word document.

    Args:
        file_path (str): The path of the Word document.
        row (int): The row index of the cell.
        column (int): The column index of the cell.
        table_number (int): The index of the table in the document.

    Returns:
        None
    """
    document = Document(file_path)
    Table = document.tables[table_number]
    run = Table.rows[row].cells[column].paragraphs[0].runs[0]
    run.font.bold = True
    document.save(file_path)


def color_alternate_rows(file_path, color_code, table_number):
    """
    Color alternate rows of a table in a Word document.

    Args:
        file_path (str): The file path of the Word document.
        color_code (str): The color code to apply to the alternate rows.
        table_number (int): The index of the table in the document.

    Returns:
        None
    """
    document = Document(file_path)
    Table = document.tables[table_number]
    num_rows = len(Table.rows)
    for row in range(1, num_rows):
        if row % 2 == 0:
            for column in range(len(Table.rows[row].cells)):
                cell_xml_element = Table.rows[row].cells[column]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), color_code)
                table_cell_properties.append(shading)
    document.save(file_path)


def color_header(file_path, color_code, table_number):
    """
    Color the header cells of a table in a Word document.

    Args:
        file_path (str): The path to the Word document.
        color_code (str): The color code to apply to the header cells.
        table_number (int): The index of the table in the document.

    Returns:
        None
    """
    document = Document(file_path)
    print(table_number)
    print(len(document.tables))
    Table = document.tables[table_number]
    for column in range(len(Table.rows[0].cells)):
        cell_xml_element = Table.rows[0].cells[column]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_code)
        table_cell_properties.append(shading)
    document.save(file_path)


def add_percent_to_pandas_df(df):
    """
    Adds a percentage sign (%) to each value in a pandas DataFrame.

    Args:
        df (pandas.DataFrame): The DataFrame to modify.

    Returns:
        pandas.DataFrame: The modified DataFrame with percentage signs added to each value.
    """
    for i in range(len(df)):
        for j in range(len(df.columns)):
            df.iloc[i, j] = str(df.iloc[i, j]) + "%"
    return df


def highlight_first_row(file_path, rgb_color, table_number):
    """
    Color alternate rows of a table in a Word document.

    Args:
        file_path (str): The file path of the Word document.
        color_code (str): The color code to apply to the alternate rows.
        table_number (int): The index of the table in the document.

    Returns:
        None
    """
    document = Document(file_path)
    Table = document.tables[table_number]
    for column in range(len(Table.rows[0].cells)):
        cell_xml_element = Table.rows[0].cells[column]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), rgb_color)
        table_cell_properties.append(shading)
        # Color the text whiite
        run = Table.rows[0].cells[column].paragraphs[0].runs[0]
        run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
    document.save(file_path)


def make_first_row_bold(file_path, table_number):
    """
    Makes the first row of a table in a Word document bold.

    Args:
        file_path (str): The path to the Word document.
        table_number (int): The index of the table in the document.

    Returns:
        None
    """
    document = Document(file_path)
    Table = document.tables[table_number]
    for column in range(len(Table.rows[0].cells)):
        run = Table.rows[0].cells[column].paragraphs[0].runs[0]
        run.font.bold = True
    document.save(file_path)


def create_table(df, file_path, shade_color):
    """
    Create a table in a Word document using the provided DataFrame.

    Args:
        df (pandas.DataFrame): The DataFrame containing the data for the table.
        file_path (str): The file path of the Word document.
        shade_color (str): The color used to shade alternate rows.

    Returns:
        None
    """
    if not check_file_exists(file_path):
        doc = docx.Document()
    else:
        doc = docx.Document(file_path)
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    print(len(doc.tables))
    for j in range(df.shape[-1]):
        table.cell(0, j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1, j).text = str(df.values[i, j])
    doc.styles['Normal'].font.name = 'Calibri'
    doc.save(file_path)
    color_alternate_rows(file_path, shade_color, -1)
    color_header(file_path, shade_color, -1)
    doc = Document(file_path)
    table = doc.tables[-1]
    table.style = 'Table Grid'
    doc.save(file_path)

# ************ END TABLE INSERTION FUNCTIONS ************ #


# ************ START PAGE BREAK ************ #


def insert_page_break(client_file_paths_list):
    """
    Inserts a page break at the end of each document in the given list of client file paths.

    Args:
        client_file_paths_list (list): A list of file paths representing the client documents.

    Returns:
        None
    """
    for i in range(len(client_file_paths_list)):
        doc = Document(client_file_paths_list[i])
        doc.add_page_break()
        doc.save(client_file_paths_list[i])

# ************ END PAGE BREAK ************ #


# ************ START MARGIN MANIPULATION ************ #


def change_margins(client_file_paths_list, top, bottom, left, right):
    """
    Change the margins of the documents in the given list of file paths.

    Args:
        client_file_paths_list (list): List of file paths of the documents.
        top (float): Top margin value in inches.
        bottom (float): Bottom margin value in inches.
        left (float): Left margin value in inches.
        right (float): Right margin value in inches.
    """
    for i in range(len(client_file_paths_list)):
        doc = docx.Document(client_file_paths_list[i])
        sections = doc.sections
        for section in sections:
            section.top_margin = top
            section.bottom_margin = bottom
            section.left_margin = left
            section.right_margin = right
        doc.save(client_file_paths_list[i])


def change_header_margins(client_file_paths_list, top, bottom, left, right):
    """
    Change the header margins of the specified Word documents.

    Args:
        client_file_paths_list (list): List of file paths of the Word documents.
        top (float): The top margin value to set.
        bottom (float): The bottom margin value to set.
        left (float): The left margin value to set.
        right (float): The right margin value to set.
    """
    for i in range(len(client_file_paths_list)):
        doc = docx.Document(client_file_paths_list[i])
        header = doc.sections[0].header
        header.top_margin = top
        header.bottom_margin = bottom
        header.left_margin = left
        header.right_margin = right
        doc.save(client_file_paths_list[i])


# ************ END MARGIN MANIPULATION ************ #


# ************ END HELPER FUNCTIONS ************ #


# ************ START FILE PATH CREATION ************ #


def generate_file_path(outer_folder_name, windows_file_path="Windows", file=None, row=None, first_name=None, last_name=None, year=None, quarter=None):
    '''    
    Generate the file path for the 401(K) preliminary report.

    Args:
        file (str): The file path.
        row (int): The row number.
        first_name (str): The first name.
        last_name (str): The last name.
        year (int): The year.
        quarter (int): The quarter.

    Returns:
        str: The file path for the 401(K) preliminary report.
    '''
    if not first_name:
        first_name = get_cell_contents(file, row, "First Name")
    if not last_name:
        last_name = get_cell_contents(file, row, "Last Name")
        print(first_name, last_name, year, quarter)

    if windows_file_path == "Windows":
        return f'{outer_folder_name}\\{last_name}, {first_name}\\{year} Q{quarter} {last_name}, {first_name} - 401(K) Preliminary Report.docx'
    else:
        return f'{outer_folder_name}/{last_name}, {first_name}/{year} Q{quarter} {last_name}, {first_name} - 401(K) Preliminary Report.docx'


def create_client_list(outer_folder_name, windows_file_path="Windows", client_excel_file=None, quarter=None, year=None):
    """
    Iterates over each row of an excel file and creates a list of file paths for each client.

    Args:
        clients_excel (str): The path to the Excel file.
        quarter (str): The quarter of the report.
        year (int): The year of the report.

    Returns:
        tuple: A tuple containing the list of client report names and the list of client names.
    """
    df = pd.read_excel(client_excel_file)

    client_names_set = set()
    for i in range(len(df)):
        current_first_name = df['First Name'][i]
        current_last_name = df['Last Name'][i]
        client_names_set.add((current_last_name, current_first_name))

    # Sort the set of names
    sorted_client_names = sorted(client_names_set)

    # Generate file paths for each sorted name
    clients_report_names = [generate_file_path(
        outer_folder_name=outer_folder_name, windows_file_path=windows_file_path, first_name=first_name, last_name=last_name, year=year, quarter=quarter)
        for last_name, first_name in sorted_client_names]

    return clients_report_names, [[last_name, first_name] for last_name, first_name in sorted_client_names]


# ************ END FILE PATH CREATION ************ #


# ************ START FILE MANAGEMENT AND DELETION ************ #


def delete_and_replace_old_files(clients_files_list):
    """
    Deletes existing files in the given list of file paths and replaces them with new blank documents.

    Args:
        clients_files_list (list): A list of file paths.

    Returns:
        None
    """
    print(clients_files_list)
    for file_path in clients_files_list:
        print(file_path)
        # Extract the directory from the file path
        directory = os.path.dirname(file_path)

        # Check if the directory exists, if not, create it
        if not os.path.exists(directory):
            os.makedirs(directory)

        # Check if the file exists
        if os.path.exists(file_path):
            # If it does, remove it
            os.remove(file_path)

        # Create a new blank document
        doc = docx.Document()
        # Save the blank document with the intended file path
        doc.save(file_path)


# ************ END FILE MANAGEMENT AND DELETION ************ #


# ************ BLANK LINE INSERTION ************ #
def add_blank_line(client_file_paths_list):
    """
    Adds a blank line to each document in the given list of client file paths.

    Args:
        client_file_paths_list (list): A list of file paths for client documents.

    Returns:
        None
    """
    for i in range(len(client_file_paths_list)):
        doc = Document(client_file_paths_list[i])
        doc.add_paragraph()
        doc.save(client_file_paths_list[i])


# ************ START TITLE INSERTION ************ #


def insert_401k_titles(clients_files_list):
    """
    Inserts titles for 401k reports into the specified client files.

    Args:
        clients_files_list (list): A list of file paths for the client files.

    Returns:
        None
    """
    add_blank_line(clients_files_list)
    for i in range(len(clients_files_list)):
        insert_paragraph_with_font_style(
            clients_files_list[i], f'{clients_files_list[i].split("/")[-1][:-5]}', 22, 'Calibri', (255, 255, 255), header=True, highlight="Blue")


# ************ END TITLE INSERTION ************ #


# ************ START IN BRIEF INSERTION ************ #


def add_in_brief(in_brief_file, client_files):
    """
    Add in brief information from a Word document to multiple client files.

    Parameters:
    - in_brief_file (str): The path of the Word document containing the in brief information.
    - client_files (list): A list of paths to the client files.

    Returns:
    None
    """
    for i in range(len(client_files)):
        copy_text_with_design_from_word_doc(in_brief_file, client_files[i])


# ************ END IN BRIEF INSERTION ************ #


# ************ START REQUIREMENTS INSERTION ************ #


def add_relevent_points_of_interest_title(client_file_paths_list):
    for i in range(len(client_file_paths_list)):
        insert_paragraph_with_font_style(
            client_file_paths_list[i], 'RELEVENT POINTS OF INTEREST', 18, 'Calibri', (76, 97, 187), header=True)


def requirements_df_to_word(df, word_file, row_color, year, quarter, header_color=None):
    """
    Convert a DataFrame to a Word document with specific formatting.

    Parameters:
    - df (pandas.DataFrame): The DataFrame to be converted.
    - word_file (str): The path to the Word document file.
    - row_color (tuple): The RGB color code for the table rows.
    - header_color (tuple, optional): The RGB color code for the table header. Defaults to None.

    Returns:
    None
    """
    print("REQUIREMENTS")
    print(df)
    new_df = df.iloc[:, -2:]
    print(new_df)
    doc = Document(word_file)
    insert_paragraph_with_font_style(
        word_file, f'{year} Q{quarter} REQUIREMENTS', 16, 'Calibri', (0, 0, 0), header=True)
    create_table(new_df, word_file, row_color)
    if header_color is not None:
        color_header(word_file, header_color, len(doc.tables) - 1)
    highlight_first_row(word_file, "#4C61BB", len(doc.tables) - 1)


def add_requirements_table(client_file_paths_list, client_names_list, requirements_excel_file, row_color, year, quarter, header_color=None):
    """
    Adds a requirements table to each client file.

    Parameters:
    - client_file_paths_list (list): A list of file paths for each client file.
    - client_names_list (list): A list of client names, where each element is a tuple containing the first and last name.
    - requirements_excel_file (str): The file path of the requirements Excel file.
    - row_color (str): The color of the rows in the table.
    - header_color (str, optional): The color of the table header. Defaults to None.

    Returns:
    None
    """
    requirements_df = pd.read_excel(requirements_excel_file)
    for i in range(len(client_file_paths_list)):

        augmented_df = extract_rows_by_name(
            requirements_df, client_names_list[i][0], client_names_list[i][1])
        shorted_df = extract_rows_by_name(
            requirements_df, "All", "All")
        if len(augmented_df) == len(shorted_df):
            insert_paragraph_with_font_style(
                client_file_paths_list[i], f"No Individual Requirements Found For {client_names_list[i][0]}, {client_names_list[i][1]}. Add Manually!!", 30, 'Calibri', (255, 255, 255), highlight="Red")
        if len(shorted_df) == 0:
            insert_paragraph_with_font_style(
                client_file_paths_list[i], f"No Requirement Found That Are To Be Assigned to All Clients. Add Manually Or Rerun The System With an Updated Excel File With Primary Requirements!!", 30, 'Calibri', (255, 255, 255), highlight="Red")

        if header_color is None:
            requirements_df_to_word(
                augmented_df, client_file_paths_list[i], row_color, year, quarter)
        else:
            requirements_df_to_word(
                augmented_df, client_file_paths_list[i], row_color, year, quarter, header_color)
        bold_first_row(client_file_paths_list[i], -1)

# ************ END REQUIREMENTS INSERTION ************ #


# ************ START GENERAL ITEMS INSERTION ************ #


def insert_general_items_bulleted_list(client_file_paths_list, general_items_file_path, client_names_list, font_size=None, font_color=None, font_style=None):
    """
    Inserts a bulleted list of general items into each client file.

    Parameters:
    - client_file_paths_list (list): A list of file paths for each client file.
    - general_items_file_path (str): The file path of the general items file.
    - client_names_list (list): A list of client names.
    - font_size (int, optional): The font size of the bulleted list. Defaults to None.
    - font_color (str, optional): The font color of the bulleted list. Defaults to None.
    - font_style (str, optional): The font style of the bulleted list. Defaults to None.
    """
    general_items_df = pd.read_excel(general_items_file_path)
    print(general_items_file_path)

    for i in range(len(client_file_paths_list)):
        insert_paragraph_with_font_style(
            client_file_paths_list[i], 'GENERAL ITEMS', 18, 'Calibri', (76, 97, 187), header=True)
        print(general_items_df)

        augmented_df = extract_rows_by_name(
            general_items_df, client_names_list[i][0], client_names_list[i][1])
        shorted_df = extract_rows_by_name(
            augmented_df, "All", "All")
        if len(augmented_df) == len(shorted_df):
            insert_paragraph_with_font_style(
                client_file_paths_list[i], f"No Individual General Items Found For {client_names_list[i][0]}, {client_names_list[i][1]}", 30, 'Calibri', (255, 255, 255), highlight="Red")
        if len(shorted_df) == 0:
            insert_paragraph_with_font_style(
                client_file_paths_list[i], f"No General Items Found For All Clients. Add Manually Or Rerun The System With an Updated Excel File With Primary Requirements!!", 30, 'Calibri', (255, 255, 255), highlight="Red")

        items_to_add = augmented_df['General Items'].tolist()
        create_numbered_list(client_file_paths_list[i], items_to_add)

        # create_bulleted_list(client_file_paths_list[i], augmented_df['General Items'].tolist(
        # ), font_size=font_size, font_color=font_color, font_style=font_style)


# ************ END GENERAL ITEMS INSERTION ************ #


# ************ START AT A GLANCE INSERTION ************ #


def create_at_a_glance_table(at_a_glance_df, word_file_path, shade_color):
    """
    Creates an "at a glance" table in a Word document.

    Parameters:
    at_a_glance_df (DataFrame): The DataFrame containing the data for the table.
    word_file_path (str): The file path of the Word document.
    shade_color (str): The color used to shade the table cells.

    Returns:
    None
    """
    print(at_a_glance_df)
    print(word_file_path)
    create_table(at_a_glance_df, word_file_path, shade_color)
    bold_first_row(word_file_path, -1)


def insert_at_a_glance(client_file_paths_list, at_a_glance_excel_file, at_a_glance_fine_print, quarter, year, shade_color):
    """
    Inserts an 'At a Glance' section into each client file.

    Parameters:
    - client_file_paths_list (list): List of file paths for each client file.
    - at_a_glance_excel_file (str): File path of the Excel file containing the 'At a Glance' data.
    - at_a_glance_fine_print (str): File path of the Word document containing the 'At a Glance' fine print.
    - quarter (int): Quarter number.
    - year (int): Year.
    - shade_color (tuple): RGB color tuple for shading the table.

    Returns:
    None
    """
    at_a_glance_df = add_percent_to_pandas_df(
        pd.read_excel(at_a_glance_excel_file))
    for i in range(len(client_file_paths_list)):
        insert_paragraph_with_font_style(
            client_file_paths_list[i], f'{year} Q{quarter} AT A GLANCE', 18, 'Calibri', (76, 97, 187), header=True)
        create_at_a_glance_table(
            at_a_glance_df, client_file_paths_list[i], shade_color)
        insert_paragraph_with_font_style(
            client_file_paths_list[i], ' ', 1, 'Calibri', (0, 0, 0))
        highlight_first_row(client_file_paths_list[i], "#4C61BB", -1)
        copy_text_with_design_from_word_doc(
            at_a_glance_fine_print, client_file_paths_list[i])


# ************ END AT A GLANCE INSERTION ************ #


# ************ START HEADER AND FOOTER INSERTION ************ #


def add_image_to_header(client_file_paths_list, image_path):
    """
    Adds an image to the header of each document in the client_file_paths_list.

    Args:
        client_file_paths_list (list): List of file paths of the client documents.
        image_path (str): Path of the image to be added to the header.

    Returns:
        None
    """
    for i in range(len(client_file_paths_list)):
        print(client_file_paths_list[i])
        doc = docx.Document(client_file_paths_list[i])
        header = doc.sections[0].header
        for para in header.paragraphs:
            del para
        paragraph = header.add_paragraph()
        run = paragraph.add_run()
        run.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run.add_picture(image_path, width=docx.shared.Inches(
            7.83), height=docx.shared.Inches(1.06))
        doc.save(client_file_paths_list[i])


def add_image_to_footer(client_file_paths_list, image_path):
    """
    Adds an image to the footer of each document in the given list of client file paths.

    Args:
        client_file_paths_list (list): A list of file paths to the client documents.
        image_path (str): The file path of the image to be added to the footer.

    Returns:
        None
    """
    for i in range(len(client_file_paths_list)):
        doc = docx.Document(client_file_paths_list[i])
        footer = doc.sections[0].footer
        for para in footer.paragraphs:
            del para
        paragraph = footer.add_paragraph()
        run = paragraph.add_run()
        run.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run.add_picture(image_path, width=docx.shared.Inches(
            7.83), height=docx.shared.Inches(1.06))
        doc.save(client_file_paths_list[i])


def main(year, quarter, outer_folder_name, windows_file_path, clients_excel_file, in_brief_file, requirements_file_path, general_items_file_path, at_a_glance_excel_file, at_a_glance_fine_print, header_image_path, footer_image_path):
    """
    Main function for creating 401k reports.

    Args:
        year (int): The year of the report.
        quarter (int): The quarter of the report.
        in_brief_file (str): The file path of the in-brief document.
        requirements_file_path (str): The file path of the requirements document.
        general_items_file_path (str): The file path of the general items document.
        at_a_glance_excel_file (str): The file path of the At-a-Glance Excel file.
        at_a_glance_fine_print (str): The fine print for the At-a-Glance section.
        header_image_path (str): The file path of the header image.
        footer_image_path (str): The file path of the footer image.
    """
    client_list = create_client_list(
        outer_folder_name, windows_file_path, clients_excel_file, quarter, year)
    client_file_paths_list = client_list[0]
    client_names = client_list[1]
    print(client_file_paths_list)
    print(client_names)
    delete_and_replace_old_files(client_file_paths_list)
    insert_401k_titles(client_file_paths_list)
    add_in_brief(in_brief_file, client_file_paths_list)
    insert_page_break(client_file_paths_list)
    add_relevent_points_of_interest_title(client_file_paths_list)
    add_requirements_table(client_file_paths_list, client_names,
                           requirements_file_path, "F0F0F0", year, quarter)  # Check to see if this is the right color
    add_blank_line(client_file_paths_list)
    insert_general_items_bulleted_list(
        client_file_paths_list, general_items_file_path, client_names)
    change_margins(client_file_paths_list, docx.shared.Inches(
        0.5), docx.shared.Inches(1.5), docx.shared.Inches(0.5), docx.shared.Inches(0.5))
    change_header_margins(client_file_paths_list, docx.shared.Inches(
        0.1), docx.shared.Inches(0.1), docx.shared.Inches(0.1), docx.shared.Inches(0.1))
    add_blank_line(client_file_paths_list)
    insert_at_a_glance(client_file_paths_list,
                       at_a_glance_excel_file, at_a_glance_fine_print, quarter, year, "F0F0F0")

    add_image_to_header(client_file_paths_list, header_image_path)
    add_image_to_footer(client_file_paths_list, footer_image_path)


# Set up the Streamlit app
st.title('401k File Processor')

# File descriptions and example file links (replace with your actual Google Drive links)
file_descriptions = {
    "Clients File": "Upload the excel file with the list of clients. [Example](https://docs.google.com/spreadsheets/d/1ZJoEZngzcO_ZJLoCQPyGxl4f-qnm2sU-j0MaAnU93X0/edit?usp=sharing)",
    "In Brief File": "Upload the In Brief document. [Example](https://docs.google.com/document/d/1FAwc02EvQdukiJXbmFL_YtJtDQtAdoIH/edit?usp=drive_link&ouid=111485210408989043988&rtpof=true&sd=true)",
    "Requirements File": "Upload the Requirements Excel file. [Example](https://docs.google.com/spreadsheets/d/1QRU7deu0Tpocsf-k0s9N8SHXOeRWEb9r/edit?usp=drive_link&ouid=111485210408989043988&rtpof=true&sd=true)",
    "General Items File": "Upload the General Items Excel file. [Example](https://docs.google.com/spreadsheets/d/1LLH8_hmP9QBJdzEPDJrDmvK_9YKERQLH/edit?usp=drive_link&ouid=111485210408989043988&rtpof=true&sd=true)",
    "At A Glance Excel File": "Upload the At A Glance Excel file. [Example](https://docs.google.com/spreadsheets/d/1CEMOWnwKhj6fCpQKB1dXSZeId6QQClZI/edit?usp=drive_link&ouid=111485210408989043988&rtpof=true&sd=true)",
    "At A Glance Fine Print File": "Upload the At A Glance Fine Print document. [Example](https://docs.google.com/document/d/14uVZM6zVs2c5OH-itORo-_zb3K2jiNLh/edit?usp=drive_link&ouid=111485210408989043988&rtpof=true&sd=true)",
    "Header Image": "Upload the Header Image (PNG or JPG) - Should be a SEFG Logo. [Example](https://drive.google.com/file/d/1C0SwsD3pnSyXllhCuhg0-C5RATk548eW/view?usp=drive_link)",
    "Footer Image": "Upload the Footer Image (PNG or JPG). [Example](https://drive.google.com/file/d/1h0V0I8bRwaV_i0uD4tVgYyyq5usBgKDa/view?usp=drive_link)"
}

# Create file uploaders with descriptions
year = st.number_input('Enter Year', min_value=2000,
                       max_value=2100, value=2021)
quarter = st.number_input('Enter Quarter', min_value=1, max_value=4, value=1)

outer_folder_name = st.text_input(
    'Enter the name of the folder to store the output files in', value="401K_Report_Output_Files")

# Define the options for the windows or mac dropdown
options = {"Windows": "Windows", "Mac": "Mac"}
# Create a selectbox for the user to choose between Windows and Mac
selected_option = st.selectbox("Select your OS:", list(options.keys()))
# Retrieve the corresponding Boolean value
windows_file_path = options[selected_option]
# Display the selected option and corresponding Boolean value (optional)
st.write(f"You selected: {selected_option}")

clients_list_file = st.file_uploader(
    file_descriptions["Clients File"], type=['xlsx'])
in_brief_file = st.file_uploader(
    file_descriptions["In Brief File"], type=['docx'])
requirements_file_path = st.file_uploader(
    file_descriptions["Requirements File"], type=['xlsx'])
general_items_file_path = st.file_uploader(
    file_descriptions["General Items File"], type=['xlsx'])
at_a_glance_excel_file = st.file_uploader(
    file_descriptions["At A Glance Excel File"], type=['xlsx'])
at_a_glance_fine_print = st.file_uploader(
    file_descriptions["At A Glance Fine Print File"], type=['docx'])
header_image_path = st.file_uploader(
    file_descriptions["Header Image"], type=['png', 'jpg'])
footer_image_path = st.file_uploader(
    file_descriptions["Footer Image"], type=['png', 'jpg'])


# Function to check missing fields
def check_missing_fields(fields):
    missing_fields = []
    for field, value in fields.items():
        if value is None or (field == "OS Selection" and value is False):
            # For OS Selection, False is a valid value, so don't count it as missing
            continue
        if not value:
            missing_fields.append(field)
    return missing_fields

# [Your existing code to set up file uploaders and other inputs]


# Button to run the main function
if st.button('Write SEFG 401(K) Reports'):
    # Dictionary of all fields with their respective values
    fields = {
        "Outer Folder Name": outer_folder_name,
        "OS Selection": windows_file_path,
        "Clients List File": clients_list_file,
        "In Brief File": in_brief_file,
        "Requirements File": requirements_file_path,
        "General Items File": general_items_file_path,
        "At A Glance Excel File": at_a_glance_excel_file,
        "At A Glance Fine Print": at_a_glance_fine_print,
        "Header Image": header_image_path,
        "Footer Image": footer_image_path
    }

    missing_fields = check_missing_fields(fields)

    if not missing_fields:
        processing_message = st.empty()
        processing_message.text("Processing... Please wait.")

        try:
            # Adjust file handling in your main function as needed
            main(
                year,
                quarter,
                outer_folder_name,
                windows_file_path,
                clients_list_file,
                in_brief_file,
                requirements_file_path,
                general_items_file_path,
                at_a_glance_excel_file,
                at_a_glance_fine_print,
                header_image_path,
                footer_image_path
            )

            processing_message.success(
                "401(K) Report Writing Completed! Your files have been uploaded to the 401K_Report_Output_Files folder.")

        except Exception as e:
            processing_message.error(f"An error occurred: {e}")

    else:
        missing_fields_message = "\n".join(
            [f"- {field}" for field in missing_fields])
        st.error(
            f"Please upload all required files. Missing:\n{missing_fields_message}")


# # Example main function call

# if __name__ == "__main__":

#     main(2021,
#          1,
#          "401K_Report_Output_Files",
#          "Mac",
#          "401k_files/Cients_List_Example.xlsx",
#          "401K_files/In Brief.docx",
#          "401k_files/test_requirements_new_2.xlsx",
#          "401k_files/test_general_items.xlsx",
#          "401k_files/at_a_glance.xlsx",
#          "401k_files/at_a_glace_fine_print.docx",
#          "401k_files/SEFG_logo.png",
#          "401k_files/Fine_Print.png"
#          )
